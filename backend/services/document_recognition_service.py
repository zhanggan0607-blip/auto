"""
证书资料内容识别服务
支持OCR识别、文本提取、数据库比对、自动更新等功能
"""
import logging
import os
import tempfile
from typing import Dict, Any, Optional, List
from datetime import datetime
from django.conf import settings
from django.db import transaction
from django.core.files.uploadedfile import InMemoryUploadedFile

from apps.enterprise.models import EnterpriseDocument, DocumentAuditLog, Enterprise
from services.aliyun_ocr_service import AliyunOCRService

logger = logging.getLogger(__name__)


class DocumentRecognitionService:
    """
    证书资料内容识别服务
    支持多种文件格式的OCR识别和内容提取
    """
    
    SUPPORTED_IMAGE_FORMATS = ['jpg', 'jpeg', 'png', 'gif', 'bmp', 'webp']
    SUPPORTED_DOC_FORMATS = ['pdf', 'doc', 'docx']
    SUPPORTED_FORMATS = SUPPORTED_IMAGE_FORMATS + SUPPORTED_DOC_FORMATS
    
    MAX_FILE_SIZE = 10 * 1024 * 1024
    
    DOCUMENT_TYPE_OCR_MAPPING = {
        'business_license': 'business_license',
        'legal_id': 'id_card',
        'qualification_cert': 'general',
        'safety_license': 'general',
        'iso_cert': 'general',
        'honor_cert': 'general',
        'other': 'general',
    }

    def __init__(self):
        self.ocr_service = AliyunOCRService()

    def validate_file(self, file_obj) -> Dict[str, Any]:
        """
        验证上传文件
        """
        if not file_obj:
            return {'valid': False, 'error': '未提供文件'}
        
        file_ext = self._get_file_extension(file_obj.name)
        if file_ext not in self.SUPPORTED_FORMATS:
            return {
                'valid': False, 
                'error': f'不支持的文件格式: {file_ext}。支持格式: {", ".join(self.SUPPORTED_FORMATS)}'
            }
        
        if hasattr(file_obj, 'size') and file_obj.size > self.MAX_FILE_SIZE:
            return {
                'valid': False,
                'error': f'文件大小超过限制 (最大 {self.MAX_FILE_SIZE // 1024 // 1024}MB)'
            }
        
        return {'valid': True, 'file_ext': file_ext}

    def _get_file_extension(self, filename: str) -> str:
        """
        获取文件扩展名
        """
        if not filename:
            return ''
        return filename.rsplit('.', 1)[-1].lower() if '.' in filename else ''

    def recognize_document(self, document: EnterpriseDocument, user=None) -> Dict[str, Any]:
        """
        识别证书内容
        """
        result = {
            'success': False,
            'content': '',
            'data': {},
            'error': None
        }
        
        try:
            document.recognition_status = 'processing'
            document.save(update_fields=['recognition_status'])
            
            if not document.file_path:
                result['error'] = '证书文件不存在'
                document.recognition_status = 'failed'
                document.recognition_error = result['error']
                document.save(update_fields=['recognition_status', 'recognition_error'])
                return result
            
            file_path = document.file_path.path
            file_ext = self._get_file_extension(document.file_path.name)
            
            if file_ext in self.SUPPORTED_IMAGE_FORMATS:
                ocr_result = self._recognize_image(file_path, document.document_type)
            elif file_ext == 'pdf':
                ocr_result = self._recognize_pdf(file_path, document.document_type)
            elif file_ext in ['doc', 'docx']:
                ocr_result = self._recognize_doc(file_path, document.document_type)
            else:
                ocr_result = {'success': False, 'error': f'不支持的文件格式: {file_ext}'}
            
            if ocr_result.get('success'):
                result['success'] = True
                result['content'] = ocr_result.get('content', '')
                result['data'] = ocr_result.get('data', {})
                
                document.extracted_content = result['content']
                document.extracted_data = result['data']
                document.recognition_status = 'completed'
                document.recognition_error = None
                document.recognition_at = datetime.now()
            else:
                result['error'] = ocr_result.get('error', '识别失败')
                document.recognition_status = 'failed'
                document.recognition_error = result['error']
            
            document.save()
            
            self._create_audit_log(
                document=document,
                action_type='recognize',
                recognition_result=result,
                is_success=result['success'],
                error_message=result.get('error'),
                user=user
            )
            
        except Exception as e:
            logger.error(f"证书识别失败: {str(e)}")
            result['error'] = str(e)
            document.recognition_status = 'failed'
            document.recognition_error = str(e)
            document.save(update_fields=['recognition_status', 'recognition_error'])
            
            self._create_audit_log(
                document=document,
                action_type='recognize',
                recognition_result=result,
                is_success=False,
                error_message=str(e),
                user=user
            )
        
        return result

    def _recognize_image(self, file_path: str, document_type: str) -> Dict[str, Any]:
        """
        识别图片文件
        """
        try:
            with open(file_path, 'rb') as f:
                image_content = f.read()
            
            ocr_type = self.DOCUMENT_TYPE_OCR_MAPPING.get(document_type, 'general')
            
            if ocr_type == 'business_license':
                return self._parse_business_license_result(
                    self.ocr_service.recognize_business_license(image_content=image_content)
                )
            elif ocr_type == 'id_card':
                return self._parse_id_card_result(
                    self.ocr_service.recognize_id_card(image_content=image_content)
                )
            else:
                return self._parse_general_result(
                    self.ocr_service.recognize_general(image_content=image_content)
                )
                
        except Exception as e:
            logger.error(f"图片识别失败: {str(e)}")
            return {'success': False, 'error': str(e)}

    def _recognize_pdf(self, file_path: str, document_type: str) -> Dict[str, Any]:
        """
        识别PDF文件
        """
        try:
            import fitz
            
            doc = fitz.open(file_path)
            content_parts = []
            
            for page_num in range(min(doc.page_count, 10)):
                page = doc[page_num]
                text = page.get_text()
                if text.strip():
                    content_parts.append(text)
                
                pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
                img_data = pix.tobytes("png")
                
                ocr_result = self.ocr_service.recognize_general(image_content=img_data)
                if ocr_result.get('success') and ocr_result.get('content'):
                    content_parts.append(ocr_result['content'])
            
            doc.close()
            
            full_content = '\n'.join(content_parts)
            
            return {
                'success': True,
                'content': full_content,
                'data': self._extract_structured_data(full_content, document_type)
            }
            
        except ImportError:
            logger.warning("PyMuPDF未安装，尝试使用pdfplumber")
            return self._recognize_pdf_with_plumber(file_path, document_type)
        except Exception as e:
            logger.error(f"PDF识别失败: {str(e)}")
            return {'success': False, 'error': str(e)}

    def _recognize_pdf_with_plumber(self, file_path: str, document_type: str) -> Dict[str, Any]:
        """
        使用pdfplumber识别PDF
        """
        try:
            import pdfplumber
            
            content_parts = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages[:10]:
                    text = page.extract_text()
                    if text:
                        content_parts.append(text)
            
            full_content = '\n'.join(content_parts)
            
            return {
                'success': True,
                'content': full_content,
                'data': self._extract_structured_data(full_content, document_type)
            }
            
        except ImportError:
            return {'success': False, 'error': 'PDF处理库未安装，请安装PyMuPDF或pdfplumber'}
        except Exception as e:
            logger.error(f"PDF识别失败: {str(e)}")
            return {'success': False, 'error': str(e)}

    def _recognize_doc(self, file_path: str, document_type: str) -> Dict[str, Any]:
        """
        识别Word文档
        """
        try:
            from docx import Document
            
            doc = Document(file_path)
            content_parts = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    content_parts.append(para.text)
            
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text for cell in row.cells)
                    content_parts.append(row_text)
            
            full_content = '\n'.join(content_parts)
            
            return {
                'success': True,
                'content': full_content,
                'data': self._extract_structured_data(full_content, document_type)
            }
            
        except ImportError:
            return {'success': False, 'error': 'python-docx未安装'}
        except Exception as e:
            logger.error(f"Word文档识别失败: {str(e)}")
            return {'success': False, 'error': str(e)}

    def _parse_business_license_result(self, ocr_result: Dict) -> Dict[str, Any]:
        """
        解析营业执照识别结果
        """
        if not ocr_result.get('success'):
            return ocr_result
        
        return {
            'success': True,
            'content': f"""
企业名称: {ocr_result.get('company_name', '')}
统一社会信用代码: {ocr_result.get('credit_code', '')}
法定代表人: {ocr_result.get('legal_person', '')}
注册资本: {ocr_result.get('registered_capital', '')}
成立日期: {ocr_result.get('establish_date', '')}
营业期限: {ocr_result.get('business_term', '')}
地址: {ocr_result.get('address', '')}
经营范围: {ocr_result.get('business_scope', '')}
""".strip(),
            'data': {
                'company_name': ocr_result.get('company_name', ''),
                'credit_code': ocr_result.get('credit_code', ''),
                'legal_person': ocr_result.get('legal_person', ''),
                'registered_capital': ocr_result.get('registered_capital', ''),
                'establish_date': ocr_result.get('establish_date', ''),
                'address': ocr_result.get('address', ''),
                'business_scope': ocr_result.get('business_scope', ''),
            }
        }

    def _parse_id_card_result(self, ocr_result: Dict) -> Dict[str, Any]:
        """
        解析身份证识别结果
        """
        if not ocr_result.get('success'):
            return ocr_result
        
        return {
            'success': True,
            'content': f"""
姓名: {ocr_result.get('name', '')}
性别: {ocr_result.get('gender', '')}
民族: {ocr_result.get('nationality', '')}
出生日期: {ocr_result.get('birth_date', '')}
住址: {ocr_result.get('address', '')}
身份证号: {ocr_result.get('id_number', '')}
""".strip(),
            'data': {
                'name': ocr_result.get('name', ''),
                'gender': ocr_result.get('gender', ''),
                'nationality': ocr_result.get('nationality', ''),
                'birth_date': ocr_result.get('birth_date', ''),
                'address': ocr_result.get('address', ''),
                'id_number': ocr_result.get('id_number', ''),
            }
        }

    def _parse_general_result(self, ocr_result: Dict) -> Dict[str, Any]:
        """
        解析通用OCR识别结果
        """
        if not ocr_result.get('success'):
            return ocr_result
        
        return {
            'success': True,
            'content': ocr_result.get('content', ''),
            'data': {
                'blocks': ocr_result.get('blocks', []),
            }
        }

    def _extract_structured_data(self, content: str, document_type: str) -> Dict[str, Any]:
        """
        从文本内容中提取结构化数据
        """
        data = {}
        
        import re
        
        cert_no_match = re.search(r'证书编号[：:]\s*([A-Za-z0-9\-]+)', content)
        if cert_no_match:
            data['certificate_no'] = cert_no_match.group(1)
        
        date_patterns = [
            (r'发证日期[：:]\s*(\d{4}[-/年]\d{1,2}[-/月]\d{1,2}日?)', 'issue_date'),
            (r'有效期至[：:]\s*(\d{4}[-/年]\d{1,2}[-/月]\d{1,2}日?)', 'expiry_date'),
            (r'有效期[：:]\s*(\d{4}[-/年]\d{1,2}[-/月]\d{1,2}日?)\s*至\s*(\d{4}[-/年]\d{1,2}[-/月]\d{1,2}日?)', 'validity_period'),
        ]
        
        for pattern, field in date_patterns:
            match = re.search(pattern, content)
            if match:
                if field == 'validity_period':
                    data['validity_start'] = match.group(1)
                    data['validity_end'] = match.group(2)
                else:
                    data[field] = match.group(1)
        
        authority_match = re.search(r'发证机关[：:]\s*(.+?)(?:\n|$)', content)
        if authority_match:
            data['issuing_authority'] = authority_match.group(1).strip()
        
        return data

    def compare_with_database(self, document: EnterpriseDocument, user=None) -> Dict[str, Any]:
        """
        将识别内容与数据库记录比对
        """
        result = {
            'success': False,
            'matches': [],
            'differences': [],
            'suggestions': [],
            'error': None
        }
        
        try:
            if not document.extracted_data:
                result['error'] = '证书尚未识别或无提取数据'
                return result
            
            enterprise = document.enterprise
            extracted = document.extracted_data
            differences = []
            suggestions = []
            
            if document.document_type == 'business_license':
                if extracted.get('credit_code') and enterprise.credit_code:
                    if extracted['credit_code'] != enterprise.credit_code:
                        differences.append({
                            'field': 'credit_code',
                            'current': enterprise.credit_code,
                            'extracted': extracted['credit_code'],
                            'suggestion': '更新统一社会信用代码'
                        })
                
                if extracted.get('legal_person') and enterprise.legal_person:
                    if extracted['legal_person'] != enterprise.legal_person:
                        differences.append({
                            'field': 'legal_person',
                            'current': enterprise.legal_person,
                            'extracted': extracted['legal_person'],
                            'suggestion': '更新法定代表人'
                        })
                
                if extracted.get('address') and enterprise.address:
                    if extracted['address'] != enterprise.address:
                        differences.append({
                            'field': 'address',
                            'current': enterprise.address,
                            'extracted': extracted['address'],
                            'suggestion': '更新企业地址'
                        })
            
            result['success'] = True
            result['differences'] = differences
            result['suggestions'] = [d['suggestion'] for d in differences]
            
            document.comparison_result = result
            document.comparison_at = datetime.now()
            document.save(update_fields=['comparison_result', 'comparison_at'])
            
            self._create_audit_log(
                document=document,
                action_type='compare',
                comparison_result=result,
                is_success=True,
                user=user
            )
            
        except Exception as e:
            logger.error(f"数据库比对失败: {str(e)}")
            result['error'] = str(e)
            
            self._create_audit_log(
                document=document,
                action_type='compare',
                comparison_result=result,
                is_success=False,
                error_message=str(e),
                user=user
            )
        
        return result

    @transaction.atomic
    def update_from_recognition(self, document: EnterpriseDocument, fields: List[str] = None, user=None) -> Dict[str, Any]:
        """
        根据识别结果更新数据库记录
        """
        result = {
            'success': False,
            'updated_fields': [],
            'error': None
        }
        
        try:
            if not document.extracted_data:
                result['error'] = '证书尚未识别或无提取数据'
                return result
            
            enterprise = document.enterprise
            extracted = document.extracted_data
            old_data = {}
            new_data = {}
            
            if fields is None:
                fields = list(extracted.keys())
            
            field_mapping = {
                'credit_code': 'credit_code',
                'legal_person': 'legal_person',
                'address': 'address',
                'business_scope': 'business_scope',
                'registered_capital': 'registered_capital',
            }
            
            for extracted_field, model_field in field_mapping.items():
                if extracted_field in fields and extracted.get(extracted_field):
                    old_value = getattr(enterprise, model_field, None)
                    new_value = extracted[extracted_field]
                    
                    if old_value != new_value:
                        old_data[model_field] = old_value
                        new_data[model_field] = new_value
                        setattr(enterprise, model_field, new_value)
                        result['updated_fields'].append(model_field)
            
            if result['updated_fields']:
                enterprise.save()
            
            if extracted.get('certificate_no') and not document.document_no:
                document.document_no = extracted['certificate_no']
                document.save(update_fields=['document_no'])
            
            result['success'] = True
            
            self._create_audit_log(
                document=document,
                action_type='update',
                old_data=old_data,
                new_data=new_data,
                update_result=result,
                is_success=True,
                user=user
            )
            
        except Exception as e:
            logger.error(f"更新记录失败: {str(e)}")
            result['error'] = str(e)
            
            self._create_audit_log(
                document=document,
                action_type='update',
                update_result=result,
                is_success=False,
                error_message=str(e),
                user=user
            )
        
        return result

    def _create_audit_log(self, document: EnterpriseDocument, action_type: str, 
                          is_success: bool = True, error_message: str = None,
                          user=None, **kwargs):
        """
        创建审计日志
        """
        try:
            DocumentAuditLog.objects.create(
                document=document,
                action_type=action_type,
                is_success=is_success,
                error_message=error_message,
                operated_by=user,
                **kwargs
            )
        except Exception as e:
            logger.error(f"创建审计日志失败: {str(e)}")

    def batch_recognize(self, document_ids: List[int], user=None) -> Dict[str, Any]:
        """
        批量识别证书
        """
        results = {
            'total': len(document_ids),
            'success': 0,
            'failed': 0,
            'details': []
        }
        
        for doc_id in document_ids:
            try:
                document = EnterpriseDocument.objects.get(id=doc_id)
                result = self.recognize_document(document, user)
                
                if result['success']:
                    results['success'] += 1
                else:
                    results['failed'] += 1
                
                results['details'].append({
                    'id': doc_id,
                    'document_name': document.document_name,
                    'success': result['success'],
                    'error': result.get('error')
                })
                
            except EnterpriseDocument.DoesNotExist:
                results['failed'] += 1
                results['details'].append({
                    'id': doc_id,
                    'success': False,
                    'error': '证书不存在'
                })
            except Exception as e:
                results['failed'] += 1
                results['details'].append({
                    'id': doc_id,
                    'success': False,
                    'error': str(e)
                })
        
        return results


class DocumentVectorService:
    """
    证书资料向量化服务
    用于AI招标公告比对
    """

    def __init__(self):
        from services.vector import embedding_service
        self.embedding_service = embedding_service

    def index_document(self, document: EnterpriseDocument) -> bool:
        """
        将证书内容索引到向量存储
        """
        if not document.is_ai_reference:
            return False

        content = document.extracted_content or ''
        if not content:
            return False

        try:
            from services.vector import enterprise_vector_store

            metadata = {
                'document_id': document.id,
                'enterprise_id': document.enterprise_id,
                'document_type': document.document_type,
                'document_name': document.document_name,
                'source': 'certificate',
            }

            success = enterprise_vector_store.add_enterprise(
                enterprise_id=f'doc_{document.id}',
                text=content,
                metadata=metadata
            )

            return success

        except Exception as e:
            logger.error(f"索引证书到向量存储失败: {str(e)}")
            return False

    def search_similar_tenders(self, document: EnterpriseDocument, top_k: int = 10) -> List[Dict]:
        """
        根据证书内容搜索相似招标公告
        """
        if not document.extracted_content:
            return []

        try:
            from services.vector import enterprise_vector_store

            results = enterprise_vector_store.search_similar(
                query_text=document.extracted_content,
                n_results=top_k
            )

            return results

        except Exception as e:
            logger.error(f"搜索相似招标公告失败: {str(e)}")
            return []
