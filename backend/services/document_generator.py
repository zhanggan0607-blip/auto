"""
文档生成服务
"""
import os
import re
import logging
from typing import Dict, Any, Optional
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from django.conf import settings
from django.utils import timezone

logger = logging.getLogger(__name__)


class DocumentGenerator:
    """
    文档生成器
    """
    def __init__(self):
        self.media_root = settings.MEDIA_ROOT
        self.template_dir = os.path.join(self.media_root, 'templates')
        self.output_dir = os.path.join(self.media_root, 'generated_docs')
        self.pdf_dir = os.path.join(self.media_root, 'generated_pdfs')
        
        for dir_path in [self.output_dir, self.pdf_dir]:
            if not os.path.exists(dir_path):
                os.makedirs(dir_path)

    def generate(self, template, tender, variables: Dict[str, Any] = None, generate_pdf: bool = True) -> Dict[str, str]:
        """
        生成文档
        """
        if variables is None:
            variables = {}
        
        default_variables = self.get_default_variables(tender)
        default_variables.update(variables)
        
        template_path = template.file_path.path
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"模板文件不存在: {template_path}")
        
        doc = Document(template_path)
        
        self.replace_variables_in_doc(doc, default_variables)
        
        timestamp = datetime.now().strftime('%Y%m%d%H%M%S')
        filename = f"{tender.id}_{timestamp}"
        
        docx_path = os.path.join(self.output_dir, f"{filename}.docx")
        doc.save(docx_path)
        
        relative_docx_path = f"generated_docs/{filename}.docx"
        
        result = {
            'docx_path': relative_docx_path,
        }
        
        if generate_pdf:
            try:
                pdf_path = self.convert_to_pdf(docx_path, filename)
                result['pdf_path'] = pdf_path
            except Exception as e:
                logger.error(f"PDF转换失败: {str(e)}")
        
        return result

    def get_default_variables(self, tender) -> Dict[str, Any]:
        """
        获取默认变量
        """
        return {
            'project_name': tender.title,
            'project_code': tender.project_code or '',
            'publish_date': str(tender.publish_date) if tender.publish_date else '',
            'deadline_date': str(tender.deadline_date) if tender.deadline_date else '',
            'region': tender.region or '',
            'industry': tender.industry or '',
            'budget': str(tender.budget) if tender.budget else '',
            'purchaser_name': tender.purchaser_name or '',
            'purchaser_contact': tender.purchaser_contact or '',
            'purchaser_phone': tender.purchaser_phone or '',
            'agency_name': tender.agency_name or '',
            'agency_contact': tender.agency_contact or '',
            'agency_phone': tender.agency_phone or '',
            'description': tender.description or '',
            'current_date': timezone.now().strftime('%Y-%m-%d'),
            'current_year': str(timezone.now().year),
        }

    def get_enterprise_variables(self, enterprise) -> Dict[str, Any]:
        """
        从企业模型获取标书模板变量（整合后的方法）
        """
        from apps.enterprise.models import EnterpriseBidConfig
        
        try:
            bid_config = enterprise.bid_config
        except EnterpriseBidConfig.DoesNotExist:
            bid_config = None
        
        primary_contact = enterprise.contacts.filter(is_primary=True).first()
        
        qualifications_data = []
        for q in enterprise.qualifications.filter(is_valid=True):
            qualifications_data.append({
                'name': q.qualification_name,
                'grade': q.grade,
                'scope': q.scope,
                'certificate_number': q.certificate_number,
                'expiry_date': q.expiry_date.strftime('%Y-%m-%d') if q.expiry_date else '',
                'issuing_authority': q.issuing_authority,
            })
        
        performances_data = []
        performances = enterprise.performances.all()
        if bid_config and bid_config.performance_years:
            from datetime import datetime, timedelta
            cutoff_date = datetime.now().date() - timedelta(days=365 * bid_config.performance_years)
            performances = performances.filter(end_date__gte=cutoff_date)
        
        for p in performances[:10]:
            performances_data.append({
                'project_name': p.project_name,
                'client_name': p.client_name,
                'contract_amount': float(p.contract_amount) if p.contract_amount else '',
                'completion_date': p.completion_date.strftime('%Y-%m-%d') if p.completion_date else '',
                'project_location': p.project_location,
                'project_scale': p.project_scale,
                'project_manager': p.project_manager,
            })
        
        variables = {
            'company_name': enterprise.name,
            'short_name': enterprise.short_name or enterprise.name,
            'credit_code': enterprise.credit_code or '',
            'legal_person': enterprise.legal_person or '',
            'registered_capital': float(enterprise.registered_capital) if enterprise.registered_capital else '',
            'establish_date': enterprise.establishment_date.strftime('%Y-%m-%d') if enterprise.establishment_date else '',
            'business_scope': enterprise.business_scope or '',
            'address': enterprise.address or '',
            'province': enterprise.province or '',
            'city': enterprise.city or '',
            'district': enterprise.district or '',
            'contact_person': primary_contact.name if primary_contact else enterprise.contact_person or '',
            'contact_phone': primary_contact.mobile if primary_contact else enterprise.contact_phone or '',
            'contact_email': primary_contact.email if primary_contact else enterprise.contact_email or '',
            'fax': enterprise.fax or '',
            'bank_name': enterprise.bank_name or '',
            'bank_account': enterprise.bank_account or '',
            'website': enterprise.website or '',
            'qualifications': qualifications_data,
            'qualifications_text': self._format_qualifications(qualifications_data),
            'performances': performances_data,
            'performances_text': self._format_performances(performances_data),
        }
        
        if bid_config:
            variables.update({
                'builder_level': bid_config.get_builder_level_display() or '',
                'builder_majors': ', '.join(bid_config.builder_majors) if bid_config.builder_majors else '',
                'accept_consortium': '是' if bid_config.accept_consortium else '否',
                'need_safety_license': '是' if bid_config.need_safety_license else '否',
                'company_certifications': ', '.join(bid_config.company_certifications) if bid_config.company_certifications else '',
            })
        
        return variables

    def _format_qualifications(self, qualifications: list) -> str:
        """
        格式化资质信息为文本
        """
        if not qualifications:
            return ''
        lines = []
        for q in qualifications:
            line = f"{q['name']}"
            if q['grade']:
                line += f" ({q['grade']})"
            lines.append(line)
        return '、'.join(lines)

    def _format_performances(self, performances: list) -> str:
        """
        格式化业绩信息为文本
        """
        if not performances:
            return ''
        lines = []
        for p in performances:
            line = f"{p['project_name']}"
            if p['contract_amount']:
                line += f"（合同金额：{p['contract_amount']}万元）"
            lines.append(line)
        return '；'.join(lines)

    def replace_variables_in_doc(self, doc, variables: Dict[str, Any]):
        """
        替换文档中的变量
        """
        for paragraph in doc.paragraphs:
            self.replace_in_paragraph(paragraph, variables)
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self.replace_in_paragraph(paragraph, variables)
        
        for section in doc.sections:
            header = section.header
            footer = section.footer
            for paragraph in header.paragraphs:
                self.replace_in_paragraph(paragraph, variables)
            for paragraph in footer.paragraphs:
                self.replace_in_paragraph(paragraph, variables)

    def replace_in_paragraph(self, paragraph, variables: Dict[str, Any]):
        """
        在段落中替换变量
        """
        for key, value in variables.items():
            placeholder = f'{{{key}}}'
            if placeholder in paragraph.text:
                for run in paragraph.runs:
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, str(value))

    def convert_to_pdf(self, docx_path: str, filename: str) -> str:
        """
        将Word文档转换为PDF
        """
        try:
            import subprocess
            
            pdf_path = os.path.join(self.pdf_dir, f"{filename}.pdf")
            
            try:
                subprocess.run([
                    'libreoffice', '--headless', '--convert-to', 'pdf',
                    '--outdir', self.pdf_dir, docx_path
                ], check=True, capture_output=True)
            except (subprocess.CalledProcessError, FileNotFoundError):
                try:
                    from docx2pdf import convert
                    convert(docx_path, pdf_path)
                except ImportError:
                    logger.warning("docx2pdf未安装，跳过PDF转换")
                    return None
            
            if os.path.exists(pdf_path):
                return f"generated_pdfs/{filename}.pdf"
            return None
        except Exception as e:
            logger.error(f"PDF转换失败: {str(e)}")
            return None

    def extract_variables_from_template(self, template_path: str) -> list:
        """
        从模板中提取变量
        """
        variables = set()
        pattern = r'\{([a-zA-Z_][a-zA-Z0-9_]*)\}'
        
        doc = Document(template_path)
        
        for paragraph in doc.paragraphs:
            matches = re.findall(pattern, paragraph.text)
            variables.update(matches)
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        matches = re.findall(pattern, paragraph.text)
                        variables.update(matches)
        
        return sorted(list(variables))


class TemplateManager:
    """
    模板管理器
    """
    def __init__(self):
        self.generator = DocumentGenerator()

    def create_template(self, name: str, template_type: str, file_path: str, user) -> Dict[str, Any]:
        """
        创建模板
        """
        variables = self.generator.extract_variables_from_template(file_path)
        
        from apps.documents.models import DocumentTemplate
        template = DocumentTemplate.objects.create(
            name=name,
            template_type=template_type,
            file_path=file_path,
            variables=variables,
            created_by=user
        )
        
        return {
            'id': template.id,
            'name': template.name,
            'variables': variables
        }

    def validate_template(self, file_path: str) -> Dict[str, Any]:
        """
        验证模板
        """
        try:
            doc = Document(file_path)
            variables = self.generator.extract_variables_from_template(file_path)
            
            return {
                'valid': True,
                'variables': variables,
                'paragraphs_count': len(doc.paragraphs),
                'tables_count': len(doc.tables)
            }
        except Exception as e:
            return {
                'valid': False,
                'error': str(e)
            }
